""" 生成zabbix 主机组资源使用率报表
    """
import os
import sys
import time
import argparse
import logging
from bisect import bisect_left
from collections import defaultdict
from datetime import datetime
from docx.shared import RGBColor
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT


def valid_date(s):
    """ 验证 argparse 传入日期格式
    """
    try:
        datetime.strptime(s, "%Y%m%d")
        return s
    except ValueError as exc:
        msg = "Not a valid date format YYYYMMDD: {0}.".format(s)
        raise argparse.ArgumentTypeError(msg) from exc


parser = argparse.ArgumentParser()
parser.add_argument("--start", required=True,
		    help="The Start Date - format YYYYMMDD", type=valid_date)
parser.add_argument("--end", required=True,
		    help="The End Date(Inclusive) - format YYYYMMDD", type=valid_date)
parser.add_argument("--output", type=str, help="output")
parser.add_argument("--topnum", type=int, default=10)
parser.set_defaults(handler=lambda args: main(args))


def convert_unit(size):
    """ 将 bytes 转换为易读的单位
        bisect_left 确定 size 在 sizes 中应插入位置 factor
        size 应除以 1024 ** factor, 加上对应的单位字符
    """
    # 各单位表示字符
    units = "BKMGTPE"
    # 各单位对应数值大小, [1024, 1024**2, 1024**3, ...]
    sizes = [1024 ** i for i in range(1, 8)]
    factor = bisect_left(sizes, size)
    return str(round(size / (1024 ** factor), 2)) + units[factor]


def get_word(api, server_hostid, path, start, end, topnum):
    """" 生成word统计报表 """

    end_timestamp = time.mktime(time.strptime(end, "%Y%m%d"))
    start_timestamp = time.mktime(time.strptime(start, "%Y%m%d"))
    document = Document()

    # 设置正文中的字体 - 微软雅黑
    document.styles["Normal"].font.name = "微软雅黑"
    document.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    paragraph_cover = document.add_paragraph("")
    paragraph_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cover = paragraph_cover.add_run("\n运管平台\n监控统计分析月报\n\n")
    run_cover.bold = True
    run_cover.font.size = Pt(36)
    run_cover.font.color.rgb = RGBColor(79, 129, 189)
    run_time = paragraph_cover.add_run("\n" + end[0:4] + "年" + end[4:6] + "月")
    run_time.bold = True
    run_time.font.size = Pt(14)
    run_time.font.color.rgb = RGBColor(79, 129, 189)

    document.add_page_break()
    # 1.汇总信息页
    run_1 = document.add_heading("", level=1).add_run("一、汇总信息")
    run_1.font.name = "微软雅黑"
    run_1.font.size = Pt(20)
    run_1._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 1.1表头
    table_total = document.add_table(
        rows=4, cols=2, style="Light List Accent 1")
    table_total.cell(0, 0).text = "统计日期"
    table_total.cell(1, 0).text = "主机组数量"
    table_total.cell(2, 0).text = "主机数量"
    table_total.cell(3, 0).text = "严重告警数量"

    # 1.2表内数据
    table_total.cell(0, 1).text = "{} - {}".format(
        time.strftime("%Y/%m/%d", time.strptime(start, "%Y%m%d")),
        time.strftime("%Y/%m/%d", time.strptime(end, "%Y%m%d")))

    # 获取主机组
    host_groups = api.hostgroup.get({
        "selectHosts": ["hostid", "name"],
        "real_hosts": True,
        "with_monitored_items": True,
        "filter": {"flags": 0}
    })

    # 主机组总数量
    groups_num = len(host_groups)
    # 主机总数量
    hosts_sum = []
    for grp in host_groups:
        hosts_sum += [host["hostid"] for host in grp["hosts"]]
        hosts_sum_num = len(set(hosts_sum))
    # 获取严重告警数量
    event_sum_num = api.event.get({
        "countOutput": True,
        "value": 1,
        "severities": [3, 4, 5],
        "time_from": start_timestamp,
        "time_till": end_timestamp
    })

    table_total.cell(1, 1).text = str(groups_num)
    table_total.cell(2, 1).text = str(hosts_sum_num)
    table_total.cell(3, 1).text = str(event_sum_num)
    document.add_page_break()

    # 2.详细统计信息页
    # 2.1 表头
    run_2 = document.add_heading("", level=1).add_run("二、详细统计信息")
    run_2.font.name = "微软雅黑"
    run_2.font.size = Pt(20)
    run_2._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 2.2 表格字段
    table_detail = document.add_table(
        rows=1, cols=12, style="Light List Accent 1")
    table_detail.autofit = False
    table_detail.cell(0, 0).width = Inches(1)
    table_detail.cell(0, 1).width = Inches(0.5)
    table_detail.cell(0, 2).width = Inches(0.75)
    table_detail.cell(0, 3).width = Inches(0.75)
    table_detail.cell(0, 4).width = Inches(0.75)
    table_detail.cell(0, 5).width = Inches(0.75)
    table_detail.cell(0, 6).width = Inches(0.75)
    table_detail.cell(0, 7).width = Inches(0.75)
    table_detail.cell(0, 8).width = Inches(0.75)
    table_detail.cell(0, 9).width = Inches(0.75)
    table_detail.cell(0, 10).width = Inches(0.75)
    table_detail.cell(0, 11).width = Inches(0.5)

    table_detail.cell(0, 0).text = "主机组"
    table_detail.cell(0, 1).text = "主机数量"
    table_detail.cell(0, 2).text = "CPU利用率(avg)"
    table_detail.cell(0, 3).text = "内存总量(avg)"
    table_detail.cell(0, 4).text = "内存利用率(max)"
    table_detail.cell(0, 5).text = "内存利用率(min)"
    table_detail.cell(0, 6).text = "内存利用率(avg)"
    table_detail.cell(0, 7).text = "磁盘总量(avg)"
    table_detail.cell(0, 8).text = "磁盘使用率(max)"
    table_detail.cell(0, 9).text = "磁盘使用率(min)"
    table_detail.cell(0, 10).text = "磁盘使用率(avg)"
    table_detail.cell(0, 11).text = "严重告警数量"

    # 2.3 获取数据
    # 2.3.1 获取 application_id
    apps = api.application.get({
        "hostids": server_hostid,
        "filter": {
            "name": ["Memory aggregation", "Filesystem aggregation", "CPU aggregation"]
        },
        "output": ["applicationid"]
    })

    # 2.3.2 获取 zabbix server 中各 hostgroup的聚合 item
    items = api.item.get({
        "hostids": server_hostid,
        "applicationids": [app["applicationid"] for app in apps],
        "output": ["name", "kay_"],
        "monitored": True,  # 已启用item
        "filter": {
            "state": 0  # 0 - normal, 1 - Not supported
        }
    })

    # 根据 Average cpu utilization 监控项确定主机组
    valid_hostgroup_names = [item["name"].split("group")[1].strip()
                             for item in items
                             if item["name"].startswith("Average cpu utilization in group")]
    host_groups = [g for g in host_groups if g["name"]
                   in valid_hostgroup_names]

    # 按主机数量排序主机组
    host_groups.sort(key=lambda x: len(x["hosts"]), reverse=True)

    # 2.3.3 设置按主机组维度统计数据的变量
    # 主机组维度按内存使用率avg数组
    memory_top_group = []
    # cpu利用率数组（主机组维度）
    cpu_top_group = []
    # 磁盘使用率（主机组维度）
    filesystem_top_group = []
    # 主机维度按内存使用率avg数组
    memory_top_host = []
    # cpu利用率数组（主机维度）
    cpu_top_host = []
    # 磁盘使用率（主机维度）
    filesystem_top_host = []

    # 2.3.4 填充表格数据
    for group in host_groups:
        group_name = group["name"]
        logging.info("正在处理数据……主机组：%s", group_name)
        logging.info("开始时间：%s", str(datetime.now()))
        host_num = len(group["hosts"])
        row = table_detail.add_row()
        row.cells[0].text = group_name
        row.cells[1].text = str(host_num)

        # 获取cpu利用率
        item_cpu_name = f"Average cpu utilization in group {group_name}"
        itemid = [item["itemid"]
                  for item in items
                  if item["name"] == item_cpu_name]
        _, _, avg_v = getcalc(
            api, itemid, start_timestamp, end_timestamp)
        row.cells[2].text = str(("%.2f" % avg_v))+"%"

        # 保留信息
        cpu_top_group.append({
            "groupname": group_name, "hostnum": host_num,
            "cpu_utilization": avg_v})

        # 获取内存总量
        item_total_memory_name = f"Total memory in group {group_name}"
        itemid = [item["itemid"]
                  for item in items
                  if item["name"] == item_total_memory_name]
        _, _, avg_v = getcalc(
            api, itemid, start_timestamp, end_timestamp)

        row.cells[3].text = convert_unit(avg_v)
        memory_dic = {"groupname": group_name,
                      "hostnum": host_num, "memory_total": avg_v}
        # 获取内存利用率
        item_utilization_memory_name = f"Memory utilization in group {group_name}"
        itemid = [item["itemid"]
                  for item in items
                  if item["name"] == item_utilization_memory_name]
        min_v, max_v, avg_v = getcalc(
            api, itemid, start_timestamp, end_timestamp)

        row.cells[4].text = str(("%.2f" % max_v))+"%"
        row.cells[5].text = str(("%.2f" % min_v))+"%"
        row.cells[6].text = str(("%.2f" % avg_v))+"%"
        memory_dic["memory_utilization"] = avg_v
        memory_top_group.append(memory_dic)

        # 获取磁盘总量
        item_total_filesystem_name = f"Total disk space in {group_name}"
        itemid = [item["itemid"]
                  for item in items
                  if item["name"] == item_total_filesystem_name]
        _, _, avg_v = getcalc(
            api, itemid, start_timestamp, end_timestamp)

        row.cells[7].text = convert_unit(avg_v)
        filesystem_dic = {"groupname": group_name,
                          "hostnum": host_num, "filesystem_total": avg_v}
        # 获取磁盘使用率
        item_utilization_filesystem_name = f"Used disk space in {group_name} (percentage)"
        itemid = [item["itemid"]
                  for item in items
                  if item["name"] == item_utilization_filesystem_name]
        min_v, max_v, avg_v = getcalc(
            api, itemid, start_timestamp, end_timestamp)

        row.cells[8].text = str(("%.2f" % max_v))+"%"
        row.cells[9].text = str(("%.2f" % min_v))+"%"
        row.cells[10].text = str(("%.2f" % avg_v))+"%"
        filesystem_dic["filesystem_utilization"] = avg_v
        filesystem_top_group.append(filesystem_dic)

        # 按主机维度处理信息，包括过滤警告，以及获取对应主机的分析数据
        host_items = api.item.get({
            "hostids": [host['hostid'] for host in group["hosts"]],
            "filter": {
                "key_": [
                    "vfs.fs.totalsize",
                    "vfs.fs.usedsize",
                    "system.cpu.util[,idle]",
                    "vm.memory.size[used]",
                    "vm.memory.size[total]"
                ],
                "state": 0
            },
            "output": ["name", "key_", "hostid"],
            "monitored": True
        })

        group_host_keys = defaultdict(dict)
        for host_item in host_items:
            host_name = [host["name"] for host in group["hosts"]
                         if host["hostid"] == host_item["hostid"]][0]
            group_host_keys[host_name][host_item["key_"]] = host_item["itemid"]

        for host_name, host_keys in group_host_keys.items():

            # 获取主机分析数据
            # 内存 used 、 total
            if host_keys.get("vm.memory.size[total]"):

                _, _, mem_avg_used = getcalc(api,
                                             host_keys["vm.memory.size[used]"],
                                             start_timestamp,
                                             end_timestamp)

                _, _, mem_avg_total = getcalc(api,
                                              host_keys["vm.memory.size[total]"],
                                              start_timestamp,
                                              end_timestamp)
                if mem_avg_total != 0:
                  # 内存 使用率
                    mem_avg_utilization = 100 * mem_avg_used / mem_avg_total
                    memory_top_host.append({"hostname": host_name,
                                            "memory_utilization": mem_avg_utilization,
                                            "memory_total": mem_avg_total,
                                            "groupname": group_name})
            # cpu 使用率
            if host_keys.get("system.cpu.util[,idle]"):
                _, _, cpu_avg_idle = getcalc(api,
                                             host_keys["system.cpu.util[,idle]"],
                                             start_timestamp,
                                             end_timestamp)
                if cpu_avg_idle != 0:
                    cpu_top_host.append({"hostname": host_name,
                                         "cpu_utilization": 100 - cpu_avg_idle,
                                         "groupname": group_name})
            # 磁盘 used 、 total
            if host_keys.get("vfs.fs.totalsize"):
                _, _, disk_avg_used = getcalc(api,
                                              host_keys["vfs.fs.usedsize"],
                                              start_timestamp,
                                              end_timestamp)
                _, _, disk_avg_total = getcalc(api,
                                               host_keys["vfs.fs.totalsize"],
                                               start_timestamp,
                                               end_timestamp)
                # 磁盘 使用率
                if disk_avg_used != 0:
                    disk_avg_utilization = 100 * disk_avg_used / disk_avg_total
                    filesystem_top_host.append({"hostname": host_name,
                                                "filesystem_utilization": disk_avg_utilization,
                                                "filesystem_total": disk_avg_used,
                                                "groupname": group_name})
        event_num = api.event.get({
            "countOutput": True,
            "hostids": [host['hostid'] for host in group["hosts"]],
            "value": 1,
            "severities": [3, 4, 5],
            "time_from": start_timestamp,
            "time_till": end_timestamp
        })
        row.cells[11].text = str(event_num)

    document.add_page_break()

    # 3. 内存使用率排行
    logging.info("3. 内存使用率排行")
    run_3 = document.add_heading("", level=1).add_run("三、内存使用率排行")
    run_3.font.name = "微软雅黑"
    run_3.font.size = Pt(20)
    run_3._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

   # 新增2级标题
    logging.info("3.1 内存使用率排行, 主机组维度")
    run_3_1_desc = document.add_heading("", level=2).add_run(
        "1、主机组维度Top" + str(topnum) + "(降序)")
    run_3_1_desc.font.name = "微软雅黑"
    run_3_1_desc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_3_1_desc.font.size = Pt(16)
    # 插入表格
    table_desc_hostgroup = document.add_table(
        rows=1, cols=4, style="Light List Accent 1")
    table_desc_hostgroup.cell(0, 0).width = Inches(6)
    table_desc_hostgroup.cell(0, 1).width = Inches(0.5)
    table_desc_hostgroup.cell(0, 2).width = Inches(0.7)
    table_desc_hostgroup.cell(0, 3).width = Inches(1.3)
    table_desc_hostgroup.cell(0, 0).text = "主机组"
    table_desc_hostgroup.cell(0, 1).text = "主机数量"
    table_desc_hostgroup.cell(0, 2).text = "内存使用率(avg)"
    table_desc_hostgroup.cell(0, 3).text = "内存总量(avg)"

    # 按内存使用率排序desc
    memory_top_group = sorted(
        memory_top_group, key=lambda x: x["memory_utilization"], reverse=True)
    for i in range(min(topnum, len(memory_top_group))):
        row = table_desc_hostgroup.add_row()
        row.cells[0].text = memory_top_group[i]["groupname"]
        row.cells[1].text = str(memory_top_group[i]["hostnum"])
        row.cells[2].text = str(
            ("%.2f" % memory_top_group[i]["memory_utilization"])) + "%"
        row.cells[3].text = str(convert_unit(
            memory_top_group[i]["memory_total"]))

    # 新增2级标题
    run_3_1_asc = document.add_heading("", level=2).add_run(
        "2、主机组维度Top" + str(topnum) + "(升序)")
    run_3_1_asc.font.name = "微软雅黑"
    run_3_1_asc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_3_1_asc.font.size = Pt(16)
    # 插入表格
    table_asc_hostgroup = document.add_table(
        rows=1, cols=4, style="Light List Accent 1")
    table_asc_hostgroup.cell(0, 0).width = Inches(6)
    table_asc_hostgroup.cell(0, 1).width = Inches(0.5)
    table_asc_hostgroup.cell(0, 2).width = Inches(0.7)
    table_asc_hostgroup.cell(0, 3).width = Inches(1.3)
    table_asc_hostgroup.cell(0, 0).text = "主机组"
    table_asc_hostgroup.cell(0, 1).text = "主机数量"
    table_asc_hostgroup.cell(0, 2).text = "内存使用率(avg)"
    table_asc_hostgroup.cell(0, 3).text = "内存总量(avg)"
    # 按内存使用率排序asc
    for i in range(len(memory_top_group)-1,
                   len(memory_top_group)-1 -
                   min(topnum, len(memory_top_group)),
                   -1):
        row = table_asc_hostgroup.add_row()
        row.cells[0].text = memory_top_group[i]["groupname"]
        row.cells[1].text = str(memory_top_group[i]["hostnum"])
        row.cells[2].text = str(
            ("%.2f" % memory_top_group[i]["memory_utilization"])) + "%"
        row.cells[3].text = str(convert_unit(
            memory_top_group[i]["memory_total"]))

    # 换页
    document.add_page_break()
    # 新增2级标题
    logging.info("3.2 内存使用率排行, 主机维度")
    run_3_2_desc = document.add_heading("", level=2).add_run(
        "3、主机维度Top" + str(topnum) + "(降序)")
    run_3_2_desc.font.name = "微软雅黑"
    run_3_2_desc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_3_2_desc.font.size = Pt(16)
    # 插入表格
    table_desc_host = document.add_table(
        rows=1, cols=4, style="Light List Accent 1")
    table_desc_host.cell(0, 0).width = Inches(3.5)
    table_desc_host.cell(0, 1).width = Inches(3.5)
    table_desc_host.cell(0, 2).width = Inches(0.7)
    table_desc_host.cell(0, 3).width = Inches(1.3)
    table_desc_host.cell(0, 0).text = "主机组"
    table_desc_host.cell(0, 1).text = "主机名"
    table_desc_host.cell(0, 2).text = "内存使用率(avg)"
    table_desc_host.cell(0, 3).text = "内存总量(avg)"
    memory_top_host = sorted(
        memory_top_host, key=lambda x: x["memory_utilization"], reverse=True)
    for i in range(min(topnum, len(memory_top_host))):
        row = table_desc_host.add_row()
        row.cells[0].text = memory_top_host[i]["groupname"]
        row.cells[1].text = memory_top_host[i]["hostname"]
        row.cells[2].text = str(
            ("%.2f" % memory_top_host[i]["memory_utilization"])) + "%"
        row.cells[3].text = str(convert_unit(
            memory_top_host[i]["memory_total"]))

    # 新增2级标题
    run_3_2_asc = document.add_heading("", level=2).add_run(
        "4、主机维度Top" + str(topnum) + "(升序)")
    run_3_2_asc.font.name = "微软雅黑"
    run_3_2_asc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_3_2_asc.font.size = Pt(16)
    # 插入表格
    table_asc_host = document.add_table(
        rows=1, cols=4, style="Light List Accent 1")
    table_asc_host.cell(0, 0).width = Inches(3.5)
    table_asc_host.cell(0, 1).width = Inches(3.5)
    table_asc_host.cell(0, 2).width = Inches(0.7)
    table_asc_host.cell(0, 3).width = Inches(1.3)
    table_asc_host.cell(0, 0).text = "主机组"
    table_asc_host.cell(0, 1).text = "主机名"
    table_asc_host.cell(0, 2).text = "内存使用率(avg)"
    table_asc_host.cell(0, 3).text = "内存总量(avg)"
    for i in range(len(memory_top_host)-1,
                   len(memory_top_host)-1 - min(topnum, len(memory_top_host)),
                   -1):
        row = table_asc_host.add_row()
        row.cells[0].text = memory_top_host[i]["groupname"]
        row.cells[1].text = memory_top_host[i]["hostname"]
        row.cells[2].text = str(
            ("%.2f" % memory_top_host[i]["memory_utilization"])) + "%"
        row.cells[3].text = str(convert_unit(
            memory_top_host[i]["memory_total"]))

    document.add_page_break()
    logging.info("4. CPU使用率排行")
    # 新增1级标题
    run_4 = document.add_heading("", level=1).add_run("四、CPU使用率排行")
    run_4.font.name = "微软雅黑"
    run_4.font.size = Pt(20)
    run_4._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 新增2级标题
    logging.info("4.1 CPU使用率排行 主机组维度")
    run_4_1_desc = document.add_heading("", level=2).add_run(
        "1、主机组维度Top" + str(topnum) + "(降序)")
    run_4_1_desc.font.name = "微软雅黑"
    run_4_1_desc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_4_1_desc.font.size = Pt(16)
    # 插入表格
    cpu_table_desc_group = document.add_table(
        rows=1, cols=3, style="Light List Accent 1")
    cpu_table_desc_group.cell(0, 0).width = Inches(6)
    cpu_table_desc_group.cell(0, 1).width = Inches(0.5)
    cpu_table_desc_group.cell(0, 2).width = Inches(0.7)
    cpu_table_desc_group.cell(0, 0).text = "主机组"
    cpu_table_desc_group.cell(0, 1).text = "主机数量"
    cpu_table_desc_group.cell(0, 2).text = "CPU使用率(avg)"
    # 按cpu使用率排序desc
    cpu_top_group = sorted(
        cpu_top_group, key=lambda x: x["cpu_utilization"], reverse=True)
    for i in range(min(topnum, len(cpu_top_group))):
        row = cpu_table_desc_group.add_row()
        row.cells[0].text = cpu_top_group[i]["groupname"]
        row.cells[1].text = str(cpu_top_group[i]["hostnum"])
        row.cells[2].text = str(
            ("%.2f" % cpu_top_group[i]["cpu_utilization"])) + "%"

    # 新增2级标题
    run_4_1_asc = document.add_heading("", level=2).add_run(
        "2、主机组维度Top" + str(topnum) + "(升序)")
    run_4_1_asc.font.name = "微软雅黑"
    run_4_1_asc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_4_1_asc.font.size = Pt(16)
    # 插入表格
    cpu_table_asc_group = document.add_table(
        rows=1, cols=3, style="Light List Accent 1")
    cpu_table_asc_group.cell(0, 0).width = Inches(6)
    cpu_table_asc_group.cell(0, 1).width = Inches(0.5)
    cpu_table_asc_group.cell(0, 2).width = Inches(0.7)
    cpu_table_asc_group.cell(0, 0).text = "主机组"
    cpu_table_asc_group.cell(0, 1).text = "主机数量"
    cpu_table_asc_group.cell(0, 2).text = "CPU使用率(avg)"
    # 按cpu使用率排序asc
    for i in range(len(cpu_top_group) - 1,
                   len(cpu_top_group) - 1 - min(topnum, len(cpu_top_group)),
                   -1):
        row = cpu_table_asc_group.add_row()
        row.cells[0].text = cpu_top_group[i]["groupname"]
        row.cells[1].text = str(cpu_top_group[i]["hostnum"])
        row.cells[2].text = str(
            ("%.2f" % cpu_top_group[i]["cpu_utilization"])) + "%"

    # 换页
    document.add_page_break()
    logging.info("4.2 主机维度Top")
    # 新增2级标题
    run_4_2_desc = document.add_heading("", level=2).add_run(
        "3、主机维度Top" + str(topnum) + "(降序)")
    run_4_2_desc.font.name = "微软雅黑"
    run_4_2_desc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_4_2_desc.font.size = Pt(16)

    # 插入表格
    cpu_table_desc_host = document.add_table(
        rows=1, cols=3, style="Light List Accent 1")
    cpu_table_desc_host.cell(0, 0).width = Inches(3.5)
    cpu_table_desc_host.cell(0, 1).width = Inches(3.5)
    cpu_table_desc_host.cell(0, 2).width = Inches(0.7)
    cpu_table_desc_host.cell(0, 0).text = "主机组"
    cpu_table_desc_host.cell(0, 1).text = "主机名"
    cpu_table_desc_host.cell(0, 2).text = "CPU平均使用率(avg)"
    cpu_top_host = sorted(
        cpu_top_host, key=lambda x: x["cpu_utilization"], reverse=True)
    for i in range(topnum):
        row = cpu_table_desc_host.add_row()
        row.cells[0].text = cpu_top_host[i]["groupname"]
        row.cells[1].text = cpu_top_host[i]["hostname"]
        row.cells[2].text = str(
            ("%.2f" % cpu_top_host[i]["cpu_utilization"])) + "%"

    # 新增2级标题
    run_4_2_asc = document.add_heading("", level=2).add_run(
        "4、主机维度Top" + str(topnum) + "(升序)")
    run_4_2_asc.font.name = "微软雅黑"
    run_4_2_asc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_4_2_asc.font.size = Pt(16)

    # 插入表格
    cpu_table_asc_host = document.add_table(
        rows=1, cols=3, style="Light List Accent 1")
    cpu_table_asc_host.cell(0, 0).width = Inches(3.5)
    cpu_table_asc_host.cell(0, 1).width = Inches(3.5)
    cpu_table_asc_host.cell(0, 2).width = Inches(0.7)
    cpu_table_asc_host.cell(0, 0).text = "主机组"
    cpu_table_asc_host.cell(0, 1).text = "主机名"
    cpu_table_asc_host.cell(0, 2).text = "CPU平均使用率(avg)"
    for i in range(len(cpu_top_host) - 1, len(cpu_top_host) - 1 - topnum, -1):
        row = cpu_table_asc_host.add_row()
        row.cells[0].text = cpu_top_host[i]["groupname"]
        row.cells[1].text = cpu_top_host[i]["hostname"]
        row.cells[2].text = str(
            ("%.2f" % cpu_top_host[i]["cpu_utilization"])) + "%"

    document.add_page_break()

    # 磁盘使用率信息表格
    logging.info("5. 磁盘使用率排行")
    run_5 = document.add_heading("", level=1).add_run("五、磁盘使用率排行")
    run_5.font.name = "微软雅黑"
    run_5.font.size = Pt(20)
    run_5._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 新增2级标题
    logging.info("5.1 磁盘使用率排行 主机组维度")
    run_5_1_desc = document.add_heading("", level=2).add_run(
        "1、主机组维度Top" + str(topnum) + "(降序)")
    run_5_1_desc.font.name = "微软雅黑"
    run_5_1_desc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_5_1_desc.font.size = Pt(16)

    # 插入表格
    filesystem_table_desc_group = document.add_table(
        rows=1, cols=4, style="Light List Accent 1")
    filesystem_table_desc_group.cell(0, 0).width = Inches(6)
    filesystem_table_desc_group.cell(0, 1).width = Inches(0.5)
    filesystem_table_desc_group.cell(0, 2).width = Inches(0.7)
    filesystem_table_desc_group.cell(0, 3).width = Inches(1.3)
    filesystem_table_desc_group.cell(0, 0).text = "主机组"
    filesystem_table_desc_group.cell(0, 1).text = "主机数量"
    filesystem_table_desc_group.cell(0, 2).text = "磁盘使用率(avg)"
    filesystem_table_desc_group.cell(0, 3).text = "磁盘总量(avg)"
    # 主机组按磁盘使用率排序desc
    filesystem_top_group = sorted(
        filesystem_top_group, key=lambda x: x["filesystem_utilization"], reverse=True)
    for i in range(min(topnum, len(filesystem_top_group))):
        row = filesystem_table_desc_group.add_row()
        row.cells[0].text = filesystem_top_group[i]["groupname"]
        row.cells[1].text = str(filesystem_top_group[i]["hostnum"])
        row.cells[2].text = str(
            ("%.2f" % filesystem_top_group[i]["filesystem_utilization"])) + "%"
        row.cells[3].text = str(convert_unit(
            filesystem_top_group[i]["filesystem_total"]))

    # 新增2级标题
    run_5_1_asc = document.add_heading("", level=2).add_run(
        "2、主机组维度Top" + str(topnum) + "(升序)")
    run_5_1_asc.font.name = "微软雅黑"
    run_5_1_asc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_5_1_asc.font.size = Pt(16)

    # 插入表格
    filesystem_table_asc_group = document.add_table(
        rows=1, cols=4, style="Light List Accent 1")
    filesystem_table_asc_group.cell(0, 0).width = Inches(6)
    filesystem_table_asc_group.cell(0, 1).width = Inches(0.5)
    filesystem_table_asc_group.cell(0, 2).width = Inches(0.7)
    filesystem_table_asc_group.cell(0, 3).width = Inches(1.3)
    filesystem_table_asc_group.cell(0, 0).text = "主机组"
    filesystem_table_asc_group.cell(0, 1).text = "主机数量"
    filesystem_table_asc_group.cell(0, 2).text = "磁盘使用率(avg)"
    filesystem_table_asc_group.cell(0, 3).text = "磁盘总量(avg)"
    # 主机组按磁盘使用率排序asc
    for i in range(len(filesystem_top_group) - 1,
                   len(filesystem_top_group) - 1 -
                   min(topnum, len(filesystem_top_group)),
                   -1):
        row = filesystem_table_asc_group.add_row()
        row.cells[0].text = filesystem_top_group[i]["groupname"]
        row.cells[1].text = str(filesystem_top_group[i]["hostnum"])
        row.cells[2].text = str(
            ("%.2f" % filesystem_top_group[i]["filesystem_utilization"])) + "%"
        row.cells[3].text = str(convert_unit(
            filesystem_top_group[i]["filesystem_total"]))

    document.add_page_break()
    logging.info("5.2 磁盘使用率排行 主机维度")
    # 新增2级标题
    run_5_2_desc = document.add_heading("", level=2).add_run(
        "3、主机维度Top" + str(topnum) + "(降序)")
    run_5_2_desc.font.name = "微软雅黑"
    run_5_2_desc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_5_2_desc.font.size = Pt(16)

    # 插入表格
    filesystem_table_desc_host = document.add_table(
        rows=1, cols=4, style="Light List Accent 1")
    filesystem_table_desc_host.cell(0, 0).width = Inches(3.5)
    filesystem_table_desc_host.cell(0, 1).width = Inches(3.5)
    filesystem_table_desc_host.cell(0, 2).width = Inches(0.7)
    filesystem_table_desc_host.cell(0, 3).width = Inches(1.3)
    filesystem_table_desc_host.cell(0, 0).text = "主机组"
    filesystem_table_desc_host.cell(0, 1).text = "主机名"
    filesystem_table_desc_host.cell(0, 2).text = "磁盘使用率(avg)"
    filesystem_table_desc_host.cell(0, 3).text = "磁盘总量(avg)"
    filesystem_top_host = sorted(
        filesystem_top_host, key=lambda x: x["filesystem_utilization"], reverse=True)
    for i in range(topnum):
        row = filesystem_table_desc_host.add_row()
        row.cells[0].text = filesystem_top_host[i]["groupname"]
        row.cells[1].text = filesystem_top_host[i]["hostname"]
        row.cells[2].text = str(
            ("%.2f" % filesystem_top_host[i]["filesystem_utilization"])) + "%"
        row.cells[3].text = str(convert_unit(
            filesystem_top_host[i]["filesystem_total"]))

    # 新增2级标题
    run_5_2_asc = document.add_heading("", level=2).add_run(
        "4、主机维度Top" + str(topnum) + "(升序)")
    run_5_2_asc.font.name = "微软雅黑"
    run_5_2_asc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_5_2_asc.font.size = Pt(16)

    # 插入表格
    filesystem_table_asc_host = document.add_table(
        rows=1, cols=4, style="Light List Accent 1")
    filesystem_table_asc_host.cell(0, 0).width = Inches(3.5)
    filesystem_table_asc_host.cell(0, 1).width = Inches(3.5)
    filesystem_table_asc_host.cell(0, 2).width = Inches(0.7)
    filesystem_table_asc_host.cell(0, 3).width = Inches(1.3)
    filesystem_table_asc_host.cell(0, 0).text = "主机组"
    filesystem_table_asc_host.cell(0, 1).text = "主机名"
    filesystem_table_asc_host.cell(0, 2).text = "磁盘使用率(avg)"
    filesystem_table_asc_host.cell(0, 3).text = "磁盘总量(avg)"
    for i in range(len(filesystem_top_host) - 1, len(filesystem_top_host) - 1 - topnum, -1):
        row = filesystem_table_asc_host.add_row()
        row.cells[0].text = filesystem_top_host[i]["groupname"]
        row.cells[1].text = filesystem_top_host[i]["hostname"]
        row.cells[2].text = str(
            ("%.2f" % filesystem_top_host[i]["filesystem_utilization"])) + "%"
        row.cells[3].text = str(convert_unit(
            filesystem_top_host[i]["filesystem_total"]))

    # 设置纸张方向为横向
    for section in document.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

    document.save(path)


def getcalc(api, itemids, time_from, time_till):
    """ 计算监控项指定时间范围内的最大值，最小值，平均值

    Args:
        api (ZAPI): zabbix api
        itemids (str | list): itemids
        time_from (str | float): 开始时间戳
        time_till (str | float): 结束时间戳

    Returns:
        tuple: (min, max, avg)
    """
    trends = api.trend.get({
        "itemids": itemids,
        "time_from": time_from,
        "time_till": time_till
    })
    if len(trends) != 0:
        values_min = []
        values_max = []
        values_avg = []
        for trend in trends:
            values_min.append(float(trend["value_min"]))
            values_max.append(float(trend["value_max"]))
            values_avg.append(float(trend["value_avg"]))
        num = len(values_avg)
        avg_value = round(sum(values_avg) / num, 2)
        min_value = min(values_min)
        max_value = max(values_max)
        return min_value, max_value, avg_value
    return 0, 0, 0


def main(args):

    zapi = args.zapi

    output_file = args.output \
        or "运管平台统计分析月报_{}_{}.docx".format(args.start, args.end)

    try:
        server_host = zapi.host.get({"filter": {"host": "Zabbix server"}})[0]

        get_word(zapi, server_host["hostid"], output_file,
                 args.start, args.end, args.topnum)

        logging.info("报表导出全部完成")
    except Exception as err:
        logging.error(err)
        sys.exit(-1)
